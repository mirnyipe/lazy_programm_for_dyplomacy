import os
import sys
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Словарь для преобразования номеров месяцев в названия
MONTH_NAMES = {
    '1': 'января', '01': 'января',
    '2': 'февраля', '02': 'февраля',
    '3': 'марта', '03': 'марта',
    '4': 'апреля', '04': 'апреля',
    '5': 'мая', '05': 'мая',
    '6': 'июня', '06': 'июня',
    '7': 'июля', '07': 'июля',
    '8': 'августа', '08': 'августа',
    '9': 'сентября', '09': 'сентября',
    '10': 'октября',
    '11': 'ноября',
    '12': 'декабря'
}


def is_part_of_document_number(context):
    """
    Проверяет, является ли число частью составного номера (например, № А3233 344/2 025)
    """
    # Паттерны, указывающие на начало номера
    patterns = [
        r'№\s*[\w-]*\d',  # №, за которым следует буква/цифра (например, № А3233)
        r'\b[А-Я]{1,2}\d{3,}',  # Буква + 3+ цифры (например, А3233)
        r'\d{3,}[/-]\d',  # Много цифр + / или - (например, 344/2)
        r'\b\d{3,}\s*\d{2,4}\b'  # Пробел между группами цифр (например, 344 025)
    ]
    context_lower = context.lower()
    for pattern in patterns:
        if re.search(pattern, context_lower):
            return True
    return False


def set_justify_alignment(doc):
    """
    Устанавливает выравнивание текста по ширине во всем документе.
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print("✅ Установлено выравнивание текста по ширине")
        return True
    except Exception as e:
        print(f"❌ Ошибка при установке выравнивания текста: {e}")
        return False


def replace_quotes(text):
    """
    Заменяет прямые двойные кавычки на типографские кавычки-лапки.
    """
    # Заменяем открывающие кавычки (кавычка в начале строки или после пробела/знака препинания)
    text = re.sub(r'(^|\s)"', r'\1«', text)
    # Заменяем закрывающие кавычки (кавычка перед пробелом/знаком препинания/концом строки)
    text = re.sub(r'"(\s|[.!?;,]|$)', r'»\1', text)
    # Для оставшихся кавычек предполагаем, что они закрывающие
    text = text.replace('"', '»')
    return text


def process_quotes(doc):
    """
    Обрабатывает замену кавычек во всем документе.
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            process_paragraph_quotes(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_quotes(paragraph)
        print("✅ Заменены прямые кавычки на типографские")
        return True
    except Exception as e:
        print(f"❌ Ошибка при обработке кавычек: {e}")
        return False


def process_paragraph_quotes(paragraph):
    """
    Обрабатывает замену кавычек в параграфе.
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return
    # Заменяем кавычки
    corrected_text = replace_quotes(full_text)
    # Если текст изменился, обновляем параграф
    if corrected_text != full_text:
        # Сохраняем форматирование первого run для применения к новому тексту
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }
        # Очищаем параграф и добавляем текст с исправлениями
        paragraph.clear()
        run = paragraph.add_run(corrected_text)
        # Применяем базовое форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def replace_special_spaces(text):
    """
    Заменяет специальные пробельные символы на обычные пробелы.
    Также сжимает множественные пробелы в один.
    """
    # Заменяем различные виды пробелов на обычный пробел
    special_spaces = [
        '\u00A0',  # неразрывный пробел
        '\u2009',  # тонкий пробел
        '\u200A',  # волосистый пробел
        '\u200B',  # нулевая ширина пробела
        '\u202F',  # узкий неразрывный пробел
        '\u205F',  # средний математический пробел
        '\u3000',  # идеографический пробел
    ]
    for space in special_spaces:
        text = text.replace(space, ' ')
    # Заменяем множественные пробелы на один
    text = re.sub(r' +', ' ', text)
    return text


def process_special_spaces(doc):
    """
    Обрабатывает замену специальных пробелов на обычные пробелы во всем документе.
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            process_paragraph_special_spaces(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_special_spaces(paragraph)
        print("✅ Заменены специальные пробелы на обычные и сжаты множественные пробелы")
        return True
    except Exception as e:
        print(f"❌ Ошибка при обработке специальных пробелов: {e}")
        return False


def process_paragraph_special_spaces(paragraph):
    """
    Обрабатывает замену специальных пробелов в параграфе.
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return
    # Заменяем специальные пробелы
    corrected_text = replace_special_spaces(full_text)
    # Если текст изменился, обновляем параграф
    if corrected_text != full_text:
        # Сохраняем форматирование первого run для применения к новому тексту
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }
        # Очищаем параграф и добавляем текст с исправлениями
        paragraph.clear()
        run = paragraph.add_run(corrected_text)
        # Применяем базовое форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def add_space_before_percent(text):
    """
    Добавляет пробел перед знаком процента, если его нет.
    """
    pattern = r'(\d+(?:[.,]\d+)?)\s*(?<!\s)%'
    return re.sub(pattern, r'\1 %', text)


def process_percent_signs(doc):
    """
    Обрабатывает добавление пробелов перед знаками процента во всем документе.
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            process_paragraph_percent_signs(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_percent_signs(paragraph)
        print("✅ Добавлены пробелы перед знаками процента")
        return True
    except Exception as e:
        print(f"❌ Ошибка при обработке знаков процента: {e}")
        return False


def process_paragraph_percent_signs(paragraph):
    """
    Обрабатывает добавление пробелов перед знаками процента в параграфе.
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return
    # Добавляем пробелы перед знаками процента
    corrected_text = add_space_before_percent(full_text)
    # Если текст изменился, обновляем параграф
    if corrected_text != full_text:
        # Сохраняем форматирование первого run для применения к новому тексту
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }
        # Очищаем параграф и добавляем текст с исправлениями
        paragraph.clear()
        run = paragraph.add_run(corrected_text)
        # Применяем базовое форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def normalize_stanitsa_abbreviations(text):
    """
    Нормализует сокращения слова "станица" к формату "ст-ца".
    """
    abbreviations = {
        r'\bстани(?:ц|цы|цей|ца|це|цам|цами|цах)\b': 'ст-ца',
        r'\bст(?:\.|\b)': 'ст-ца',
        r'\bста(?:н|н\.)\b': 'ст-ца',
        r'\bстц\b': 'ст-ца',
        r'\bстани\b': 'ст-ца',
    }
    for pattern, replacement in abbreviations.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text


def process_stanitsa_abbreviations(doc):
    """
    Обрабатывает нормализацию сокращений "станица" во всем документе.
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            process_paragraph_stanitsa_abbreviations(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_stanitsa_abbreviations(paragraph)
        print("✅ Нормализованы сокращения слова 'станица'")
        return True
    except Exception as e:
        print(f"❌ Ошибка при нормализации сокращений 'станица': {e}")
        return False


def process_paragraph_stanitsa_abbreviations(paragraph):
    """
    Обрабатывает нормализацию сокращений "станица" в параграфе.
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return
    # Нормализуем сокращения
    normalized_text = normalize_stanitsa_abbreviations(full_text)
    # Если текст изменился, обновляем параграф
    if normalized_text != full_text:
        # Сохраняем форматирование первого run для применения к новому тексту
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }
        # Очищаем параграф и добавляем нормализованный текст
        paragraph.clear()
        run = paragraph.add_run(normalized_text)
        # Применяем базовое форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def normalize_dates_in_text(text):
    """
    Преобразует даты в тексте к формату "12 марта 2024 г."
    """

    def replace_dd_mm_yyyy(match):
        day, month, year = match.groups()
        try:
            month_key = str(int(month))
            month_name = MONTH_NAMES.get(month_key)
            if not month_name:
                return match.group()
            if len(year) == 2:
                year = f"20{year}" if int(year) < 30 else f"19{year}"
            return f"{int(day)} {month_name} {year} г."
        except (ValueError, KeyError):
            return match.group()

    def replace_yyyy_mm_dd(match):
        year, month, day = match.groups()
        try:
            month_key = str(int(month))
            month_name = MONTH_NAMES.get(month_key)
            if not month_name:
                return match.group()
            return f"{int(day)} {month_name} {year} г."
        except (ValueError, KeyError):
            return match.group()

    def replace_day_month_year(match):
        day, month, year = match.groups()
        return f"{day} {month} {year} г."

    def replace_day_month_abbr_year(match):
        day, month_abbr, year = match.groups()
        month_full = {
            'янв': 'января', 'фев': 'февраля', 'мар': 'марта',
            'апр': 'апреля', 'май': 'мая', 'июн': 'июня',
            'июл': 'июля', 'авг': 'августа', 'сен': 'сентября',
            'окт': 'октября', 'ноя': 'ноября', 'дек': 'декабря'
        }.get(month_abbr.lower(), month_abbr)
        return f"{day} {month_full} {year} г."

    def replace_day_month_year_g_dot(match):
        day, month, year = match.groups()
        return f"{day} {month} {year} г."

    def replace_day_month_abbr_year_g_dot(match):
        day, month_abbr, year = match.groups()
        month_full = {
            'янв': 'января', 'фев': 'февраля', 'мар': 'марта',
            'апр': 'апреля', 'май': 'мая', 'июн': 'июня',
            'июл': 'июля', 'авг': 'августа', 'сен': 'сентября',
            'окт': 'октября', 'ноя': 'ноября', 'дек': 'декабря'
        }.get(month_abbr.lower(), month_abbr)
        return f"{day} {month_full} {year} г."

    text = re.sub(r'\b(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})\b', replace_dd_mm_yyyy, text)
    text = re.sub(r'\b(\d{4})[./-](\d{1,2})[./-](\d{1,2})\b', replace_yyyy_mm_dd, text)
    text = re.sub(r'\b(\d{1,2})\s+(янв|фев|мар|апр|май|июн|июл|авг|сен|окт|ноя|дек)[.]\s*(\d{4})\b',
                  replace_day_month_abbr_year, text)
    text = re.sub(r'\b(\d{1,2})\s+(янв|фев|мар|апр|май|июн|июл|авг|сен|окт|ноя|дек)[.]\s*(\d{4})\s+г\.\b',
                  replace_day_month_abbr_year_g_dot, text)
    text = re.sub(
        r'\b(\d{1,2})\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+(\d{4})\s+г\.\b',
        replace_day_month_year_g_dot, text)
    text = re.sub(
        r'\b(\d{1,2})\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+(\d{4})\b(?!\s*г)',
        replace_day_month_year, text)
    return text


def normalize_dates(doc):
    """
    Нормализует даты во всем документе
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            normalize_paragraph_dates(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        normalize_paragraph_dates(paragraph)
        print("✅ Даты нормализованы")
        return True
    except Exception as e:
        print(f"❌ Ошибка при нормализации дат: {e}")
        return False


def normalize_paragraph_dates(paragraph):
    """
    Нормализует даты в параграфе
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return
    # Нормализуем даты в тексте
    normalized_text = normalize_dates_in_text(full_text)
    # Если текст изменился, обновляем параграф
    if normalized_text != full_text:
        # Сохраняем форматирование первого run для применения к новому тексту
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }
        # Очищаем параграф и добавляем нормализованный текст
        paragraph.clear()
        run = paragraph.add_run(normalized_text)
        # Применяем базовое форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def convert_decimal_separator_in_text(text):
    """
    Преобразует десятичные разделители в числах с точки на запятую
    """

    def replace_decimal_point(match):
        return match.group().replace('.', ',')

    pattern = r'(?<!\d[.,])\b\d+\.\d+\b(?![.,]\d)'
    text = re.sub(pattern, replace_decimal_point, text)
    return text


def process_decimal_separators(doc):
    """
    Обрабатывает замену десятичных разделителей во всем документе
    """
    try:
        # Обрабатываем параграфы в основном документе
        for paragraph in doc.paragraphs:
            process_paragraph_decimal_separators(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_decimal_separators(paragraph)
        print("✅ Заменены десятичные разделители (точка → запятая)")
        return True
    except Exception as e:
        print(f"❌ Ошибка при замене десятичных разделителей: {e}")
        return False


def process_paragraph_decimal_separators(paragraph):
    """
    Обрабатывает замену десятичных разделителей в параграфе
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return
    # Заменяем десятичные разделители
    converted_text = convert_decimal_separator_in_text(full_text)
    # Если текст изменился, обновляем параграф
    if converted_text != full_text:
        # Сохраняем форматирование первого run для применения к новому тексту
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }
        # Очищаем параграф и добавляем текст с заменами
        paragraph.clear()
        run = paragraph.add_run(converted_text)
        # Применяем базовое форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def is_likely_date(text):
    """
    Проверяет, является ли текст похожим на дату
    """
    date_patterns = [
        r'\d{1,2}[./-]\d{1,2}[./-]\d{2,4}',
        r'\d{4}[./-]\d{1,2}[./-]\d{1,2}',
        r'\d{1,2}\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}',
        r'\d{1,2}\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}\s*г\.?',
        r'\d{1,2}\s+(янв|фев|мар|апр|май|июн|июл|авг|сен|окт|ноя|дек)[.]\s*\d{4}',
    ]
    text_lower = text.lower().strip()
    for pattern in date_patterns:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return True
    return False


def contains_month_word(text):
    """
    Проверяет, содержит ли текст название месяца после числа
    """
    month_patterns = [
        r'\d+\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)',
        r'\d+\s+(янв|фев|мар|апр|май|июн|июл|авг|сен|окт|ноя|дек)[.]'
    ]
    text_lower = text.lower()
    for pattern in month_patterns:
        if re.search(pattern, text_lower):
            return True
    return False


def contains_year_word_to_exclude(text, number_text):
    """
    Проверяет, содержит ли текст слово "год" или "г." непосредственно после указанного числа.
    """
    exclude_year_patterns = [
        rf'\b{re.escape(number_text)}\s+год\b',
        rf'\b{re.escape(number_text)}\s*г\.',
    ]
    text_lower = text.lower()
    for pattern in exclude_year_patterns:
        if re.search(pattern, text_lower):
            return True
    return False


def make_numbers_bold(doc):
    """
    Выделяет жирным все числа (кроме дат и чисел с "год" и "г."), но не выделяет числа в составе номеров (№ А3233 344/2 025)
    """
    try:
        number_patterns = [
            r'[+-]?\d{1,3}(?:\s\d{3})*(?:[.,]\d+)?',
            r'[+-]?\d{1,3}(?:\u2009\d{3})*(?:[.,]\d+)?',
            r'[+-]?\d{1,3}(?:,\d{3})*(?:[.,]\d+)?',
            r'[+-]?\d{1,3}(?:\.\d{3})*(?:[.,]\d+)?',
            r'[+-]?\d+(?:[.,]\d+)?',
        ]
        for paragraph in doc.paragraphs:
            process_paragraph_numbers(paragraph, number_patterns)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_numbers(paragraph, number_patterns)
        print("✅ Числа выделены жирным (даты, 'год' и номера дел исключены)")
        return True
    except Exception as e:
        print(f"❌ Ошибка при выделении чисел: {e}")
        return False


def process_paragraph_numbers(paragraph, number_patterns):
    """
    Обрабатывает числа в параграфе
    """
    runs_text = [run.text for run in paragraph.runs]
    full_text = ''.join(runs_text)
    if not full_text.strip():
        return

    numbers_found = []
    for pattern in number_patterns:
        for match in re.finditer(pattern, full_text):
            number_text = match.group()
            start_pos = match.start()
            end_pos = match.end()
            # Расширенный контекст
            context_start = max(0, start_pos - 50)
            context_end = min(len(full_text), end_pos + 50)
            context = full_text[context_start:context_end]

            if (is_likely_date(context) or
                    contains_month_word(context) or
                    contains_year_word_to_exclude(context, number_text) or
                    is_part_of_document_number(context)):
                continue  # Пропускаем

            numbers_found.append({
                'text': number_text,
                'start': start_pos,
                'end': end_pos
            })

    # Удаление пересекающихся совпадений
    if numbers_found:
        numbers_found.sort(key=lambda x: x['start'])
        filtered = []
        for num in numbers_found:
            if not any(n['start'] < num['end'] and n['end'] > num['start'] for n in filtered):
                filtered.append(num)
        numbers_found = filtered

    if not numbers_found:
        return

    paragraph.clear()
    last_pos = 0
    for num in numbers_found:
        if num['start'] > last_pos:
            before = full_text[last_pos:num['start']]
            run = paragraph.add_run(before)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = False
        bold_run = paragraph.add_run(num['text'])
        bold_run.bold = True
        bold_run.font.name = 'Times New Roman'
        bold_run.font.size = Pt(14)
        last_pos = num['end']
    if last_pos < len(full_text):
        after = full_text[last_pos:]
        run = paragraph.add_run(after)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = False


def reset_text_formatting_except_bold(doc):
    """
    Сбрасывает все форматирование текста, кроме жирного выделения
    """
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                is_bold = run.bold
                run.font.name = None
                run.font.size = None
                run.font.bold = None
                run.font.italic = None
                run.font.underline = None
                run.font.color.rgb = None
                if is_bold is not None:
                    run.bold = is_bold
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            is_bold = run.bold
                            run.font.name = None
                            run.font.size = None
                            run.font.bold = None
                            run.font.italic = None
                            run.font.underline = None
                            run.font.color.rgb = None
                            if is_bold is not None:
                                run.bold = is_bold
        print("✅ Форматирование текста сброшено (сохранено только жирное выделение)")
        return True
    except Exception as e:
        print(f"❌ Ошибка при сбросе форматирования: {e}")
        return False


def apply_uniform_formatting(doc):
    """
    Применяет единый стиль ко всему документу
    """
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            pf = paragraph.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)
                        pf = paragraph.paragraph_format
                        pf.line_spacing = 1.5
        print("✅ Установлен единый стиль: Times New Roman, 14pt, интервал 1.5")
        return True
    except Exception as e:
        print(f"❌ Ошибка при применении единого стиля: {e}")
        return False


def format_thousands_separator_in_text(text):
    """
    Форматирует числа с разделителями тысяч (пробелы)
    Пример: 1000 -> 1 000, 2500000 -> 2 500 000, 12345.67 -> 12 345,67
    """

    def format_match(match):
        sign = match.group(1)  # знак (+ или -)
        integer_part = match.group(2)  # целая часть
        decimal_part = match.group(3)  # десятичная часть с разделителем

        # Форматируем целую часть: разбиваем на группы по 3 цифры
        integer_rev = integer_part[::-1]
        chunks = [integer_rev[i:i + 3] for i in range(0, len(integer_rev), 3)]
        formatted_integer = ' '.join(chunks)[::-1]

        return sign + formatted_integer + decimal_part

    # Шаблон для чисел:
    #   [+-]? - необязательный знак
    #   \d{1,3} - от 1 до 3 цифр
    #   (?:\d{3})* - группы по 3 цифры (0 или более)
    #   (?:[.,]\d+)? - необязательная десятичная часть
    pattern = r'\b([+-]?)(\d{1,3}(?:\d{3})*)([.,]?\d*)\b'
    return re.sub(pattern, format_match, text)


def process_thousands_separator(doc):
    """
    Форматирует числа с разделителями тысяч во всем документе
    """
    try:
        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            process_paragraph_thousands_separator(paragraph)
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph_thousands_separator(paragraph)
        print("✅ Числа отформатированы с разделителями тысяч (пробел)")
        return True
    except Exception as e:
        print(f"❌ Ошибка при форматировании разделителей тысяч: {e}")
        return False


def process_paragraph_thousands_separator(paragraph):
    """
    Форматирует числа с разделителями тысяч в параграфе
    """
    # Собираем весь текст параграфа
    full_text = ''.join([run.text for run in paragraph.runs])
    if not full_text.strip():
        return

    # Форматируем числа
    formatted_text = format_thousands_separator_in_text(full_text)

    # Если текст изменился, обновляем параграф
    if formatted_text != full_text:
        # Сохраняем форматирование первого run
        first_run_format = {}
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline
            }

        # Очищаем параграф и добавляем отформатированный текст
        paragraph.clear()
        run = paragraph.add_run(formatted_text)

        # Восстанавливаем форматирование
        run.font.name = first_run_format.get('font_name', 'Times New Roman')
        run.font.size = first_run_format.get('font_size', Pt(14))
        if first_run_format.get('bold') is not None:
            run.bold = first_run_format.get('bold')
        if first_run_format.get('italic') is not None:
            run.italic = first_run_format.get('italic')
        if first_run_format.get('underline') is not None:
            run.underline = first_run_format.get('underline')


def set_document_margins(doc_path):
    """
    Основная функция обработки документа
    """
    try:
        if not os.path.exists(doc_path):
            print(f"❌ Файл не найден: {doc_path}")
            return False
        doc = Document(doc_path)
        for i, section in enumerate(doc.sections):
            print(f"Обрабатываем секцию {i + 1}")
            section.top_margin = Cm(1.0)
            section.right_margin = Cm(1.5)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.5)

        if not reset_text_formatting_except_bold(doc):
            return False
        if not apply_uniform_formatting(doc):
            return False
        if not process_special_spaces(doc):
            return False
        if not process_quotes(doc):
            return False
        if not normalize_dates(doc):
            return False
        if not process_decimal_separators(doc):
            return False
        if not process_percent_signs(doc):
            return False
        if not process_stanitsa_abbreviations(doc):
            return False
        if not process_thousands_separator(doc):  # Новый шаг форматирования тысяч
            return False
        if not make_numbers_bold(doc):
            return False
        if not set_justify_alignment(doc):
            return False

        name, ext = os.path.splitext(doc_path)
        output_path = f"{name}_formatted{ext}"
        doc.save(output_path)
        print(f"✅ Успешно! Документ сохранён как: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        return False


def main():
    print("=== Редактор Word документов ===")
    print("Выполняемые действия:")
    print("1. Установка полей: Верх=1см, Право=1.5см, Низ=1см, Лево=1.5см")
    print("2. Сброс форматирования (сохранено только жирное выделение)")
    print("3. Установка стиля: Times New Roman, 14pt, интервал 1.5")
    print("4. Замена специальных пробелов на обычные")
    print("5. Замена прямых кавычек на типографские")
    print("6. Нормализация дат")
    print("7. Замена десятичных разделителей (точка → запятая)")
    print("8. Добавление пробелов перед знаками процента")
    print("9. Нормализация сокращений 'станица'")
    print("10. Форматирование разделителей тысяч (1 000)")
    print("11. Выделение чисел жирным (даты, 'год' и номера дел исключены)")
    print("12. Выравнивание по ширине")
    print("-" * 65)

    if len(sys.argv) > 1:
        file_path = sys.argv[1].strip('"\'')
    else:
        file_path = input("Введите путь к .docx файлу: ").strip().strip('"\'')
    if not file_path.lower().endswith('.docx'):
        print("⚠️ Файл должен иметь расширение .docx")
        return

    success = set_document_margins(file_path)
    if not success:
        print("❌ Обработка завершена с ошибками")
    else:
        print("🎉 Обработка завершена успешно!")


if __name__ == "__main__":
    main()
